import patterns
import init_paths
import re
import PyPDF2
import docx

class ExtractData:
    def __init__(self, name) -> None:
        self.name = name
        self.pathstruct = init_paths(self.name)
        if self.name=='service_recommendation':
            self.data_patterns = patterns.get(
                'service_recommendation'
            ).get('data_patterns')
            self.labels_patterns = patterns.get(
                'service_recommendation'
            ).get('label_patterns')
        elif self.name=='credentials_recommendation':
            self.data_patterns = patterns.get(
                'credentials_recommendation'
            ).get('data_patterns')
            self.labels_patterns = patterns.get(
                'credentials_recommendation'
            ).get('label_patterns')
        else:
            raise Exception
    
    def extract(self, file_type):
        foi = self.pathstruct.create_paths(file_type)
        
        if file_type == 'pdf':
            data = []
            labels = []
            pattern = patterns.patterns.get('service_recommendation')
            # to be replaced by self.data_patterns and self.label_patterns
            for file in foi:
                reader = PyPDF2.PdfReader(file)
                if reader.is_encrypted:
                    try:
                        reader.decrypt('client')
                        data.extend(self.pdf_extract(reader, pattern.get('data_patterns')))
                        labels.extend(self.pdf_extract(reader, pattern.get('labels_patterns')))
                    except PyPDF2.errors.FileNotDecryptedError:
                        print(f'password for {file.stem} was incorrect')
                else:
                    data.extend(self.pdf_extract(reader, pattern.get('data_patterns')))
                    labels.extend(self.pdf_extract(reader, pattern.get('labels_patterns')))
        elif file_type == 'docx':
            names = []
            names_and_files = []
            labels = []
            pattern = patterns.patterns.get('credentials_recommendation')
            for file in foi:
                names_res = self.docx_extract(file, pattern.get('data_patterns'))
                names.extend(names_res)
                names_and_files.append([names_res, file])
                # labels.extend(self.docx_extract(reader, pattern.get('labels_patterns')))
            return names, names_and_files

    def pdf_extract(reader, patterns):
            tmp_data = []
            tmp_text = ' '
            for j in range(len(reader.pages)):
                tmp_text += reader.pages[j].extract_text()
            tmp_text = tmp_text.strip().replace('\n', ' ')
            flags = re.IGNORECASE | re.DOTALL

            def process_patterns(patterns, text):
                for pattern in patterns:
                    matches = re.findall(pattern, text, flags)
                    if matches:
                        return matches
                return None

            matched = False
            for key, pattern_set in patterns.items():
                if isinstance(pattern_set, list):
                    for sub_patterns in pattern_set:
                        if isinstance(sub_patterns, list):
                            matches = process_patterns(sub_patterns, tmp_text)
                        else:
                            matches = re.findall(sub_patterns, tmp_text, flags)
                        if matches:
                            tmp_data.extend(matches)
                            matched = True
                            break
                elif isinstance(pattern_set, str):
                    if key == 'client care plan 1':
                        ignore_terms = ['Bathing \\(bath\\)', 'Dressing \\(lower\\)', 
                                        'Wheelchair', 'Hair Care', 'Skin Care', 'Falls', 'Morning Meds']
                        ignore_pattern = r'\b(?:{})\b'.format('|'.join(ignore_terms))
                        clean_service_text = re.sub(ignore_pattern, '', tmp_text, flags)
                        matches = re.findall(pattern_set, clean_service_text, flags)
                    else:
                        matches = re.findall(pattern_set, tmp_text, flags)
                    if matches:
                        tmp_data.extend(matches)
                        matched = True
                        break
                else:
                    raise Exception('Not possible!')

            if not matched:
                tmp_data.append(None)

            return tmp_data
    
    def docx_extract(file, patterns):
        tmp_data = []
        tmp_text = ' '
        doc = docx.Document(file)
        for para in doc.paragraphs:
            tmp_data.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Avoid duplicate text entries by checking for cell contents in the text list
                    if cell.text not in tmp_data:
                        tmp_data.append(cell.text)
            
        tmp_text = ' '.join(tmp_data)
        matches = re.findall(patterns, tmp_text)
        extracted_names = [''.join(match).strip() for match in matches]
        extracted_names = set(extracted_names)
        return extracted_names
        
        




       
# def tf_extract_info(self, text):
#     """
#     a helper function to map a preprocess object method to 
#     extract info needed.
#     """
#     def extract_info_py(text):
#         extracted = preo.extract_info(text.numpy().decode('utf-8'))
#         return (
#         tensorflow.constant(extracted['names']),
#         tensorflow.constant(extracted['emails']),
#         tensorflow.constant(extracted['phones']),
#         )
#     names, emails, phones = tensorflow.py_function(
#         extract_info_py, [text], [
#             tensorflow.string,
#             tensorflow.string,
#             tensorflow.string,
#         ]
#     )
#     return names, emails, phones
    